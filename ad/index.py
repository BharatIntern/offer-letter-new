import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx2pdf import convert
import yagmail
from datetime import datetime

# Step 1: Read data from Excel file
data = pd.read_excel('data.xlsx')

if 'Email_Status' not in data.columns:
    data['Email_Status'] = ''
if 'Email_Timestamp' not in data.columns:
    data['Email_Timestamp'] = ''


# Step 2: Loop through each row in the Excel file
for i in range(len(data)):
    # check if email has already been sent
    if data['Email_Status'][i] == 'Sent':
        print('Email already sent to', data['Name'][i], '<', data['Email'][i], '>')
        continue

    # load the Word document template
    document = Document('template.docx')

    for p in document.paragraphs:
        if '<<name>>' in p.text:
            tag_text = p.text
            p.clear()  # remove the tag
            run = p.add_run(tag_text.replace('<<name>>', data['Name'][i]))
            tag_font = run.font
            tag_font.color.rgb = RGBColor(149, 74, 0)  # set font color
            tag_font.bold = True  # set bold
            tag_font.name = 'Comic Sans MS'  # set font style
            tag_font.size = Pt(19)  # set font size

    # save the updated document as a new file
    updated_template_name = '{}.docx'.format(data['Name'][i].replace(' ', '_'))
    document.save(updated_template_name)

    # convert the Word document to PDF and save in attachments folder
    pdf_file = os.path.join('attachments', '{}.pdf'.format(data['Name'][i].replace(' ', '_')))
    convert(updated_template_name, pdf_file)

    # create email
    html_content = """ <p><strong>Congratulations  😍🎉, you are selected for Virtual Internship Program at Bharat Intern.


Upload your welcome letter on our LinkedIn page and tag us  <a href="https://www.linkedin.com/company/bharat-intern/">BharatIntern</a>


Task List - <a href="https://drive.google.com/file/d/1F98hMgM1vIq4yhE-xooQ7Qw8jh9SEeUV/view?usp=sharing">PDF</a> 

Some important points to remember during your internship tenure.


Task submission procedure -


1. Atleast 2 tasks(Check PDF) are mandatory to be completed for the completion of internship and to become eligible for the certification.

2. Maintain a separate GitHub repository for all the tasks and share the link of the GitHub repo in the task submission form(it will be given later on your email and telegram).

3. Make a short video (under 1-minute) or 3-4 photos of the code of the completed task and share it on LinkedIn and tag  <a href="https://www.linkedin.com/company/bharat-intern/">BharatIntern</a>

The timeline will be this way!


10 March  - Internship started


1 April  - Task submission form will be sent to all the interns.


10 April - Last date of task submission.


Till 15 April - The internship completion certificate will be sent to deserving candidates.

For any query, you can contact us on Telegram-<a href="https://t.me/bharatintern">BharatIntern</a>

Best Regards,
Team BharatIntern.
</strong></p>"""

    receiver = data['Email'][i]

    with yagmail.SMTP('bharatintern.info', 'lbirzenvzklnfovs') as yag:
        yag.send(
            to=receiver,
            subject='Offer Letter and Task(Bharat Intern)',
            contents=[html_content.format(data['Name'][i]), pdf_file],
        )

    print('Email sent to', data['Name'][i], '<', data['Email'][i], '>')

    # update email status and timestamp in Excel file
    data.at[i, 'Email_Status'] = 'Sent'
    data.at[i, 'Email_Timestamp'] = datetime.now()
    data.to_excel('data.xlsx', index=False)
    # remove the updated template file
    os.remove(updated_template_name)
    print(f"Mail No:{i+1}")

# save the updated Excel file
